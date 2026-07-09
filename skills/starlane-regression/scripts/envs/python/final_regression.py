"""Final regression execution and table rendering for the Python env (library)."""

from __future__ import annotations

import csv
import math
import os
from pathlib import Path

import pandas as pd

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from common import (
    RegressionArgs,
    apply_spec_condition,
    build_model_plan,
    copy_source_to_dir,
    encode_panel_if_needed,
    ensure_columns,
    format_coef,
    make_base_sample,
    parse_discrete_values,
    parse_rob_vars,
    prepare_regression_data,
    read_data,
    run_spec,
    spec_required_columns,
    split_words,
    stars_for_result,
)


SECTION_TITLES = {
    "baseline": "基准回归",
    "rob_alt_x": "稳健性检验-替换X",
    "rob_alt_y": "稳健性检验-替换变量",
    "rob_ln_x": "稳健性检验-X取对数",
    "rob_ln_y": "稳健性检验-Y取对数",
    "rob_lag": "稳健性检验-滞后期",
    "rob_year": "稳健性检验-时间窗口",
    "iv_stage1": "工具变量-一阶段",
    "iv_stage2": "工具变量-二阶段 2SLS",
    "med": "中介机制",
    "mod": "异质性分析-调节效应检验",
}


def section_title(section: str) -> str:
    if section in SECTION_TITLES:
        return SECTION_TITLES[section]
    if section.startswith("het_disc_"):
        return f"Heterogeneity: {section.removeprefix('het_disc_')}"
    return section.replace("_", " ").title()


def doc_section_key(section: str) -> str:
    if section in ("baseline_nocv", "baseline_cv"):
        return "baseline"
    if section.startswith("med_"):
        return "med"
    if section.startswith("mod_"):
        return "mod"
    return section


def model_label(row: dict[str, str], idx: int) -> str:
    if row["condition"]:
        return row["condition"]
    return f"({idx})"


def format_decimal(value: str | float, digits: int = 3) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return ""
    if not math.isfinite(number):
        return ""
    return f"{number:.{digits}f}"


def format_n(value: str | int) -> str:
    try:
        number = int(float(value))
    except (TypeError, ValueError):
        return ""
    return f"{number:,}"


def stars_from_p(p_value: object) -> str:
    try:
        p = float(p_value)
    except (TypeError, ValueError):
        return ""
    if not math.isfinite(p):
        return ""
    if p < 0.01:
        return "***"
    if p < 0.05:
        return "**"
    if p < 0.1:
        return "*"
    return ""


def coef_text(row: dict[str, str], variable: str) -> str:
    coefficients = row.get("coefficients", {})
    if not isinstance(coefficients, dict) or variable not in coefficients:
        return ""
    if variable == row["target"]:
        stars = row["stars"]
    else:
        p_values = row.get("p_values", {})
        stars = stars_from_p(p_values.get(variable)) if isinstance(p_values, dict) else ""
    return f"{format_decimal(coefficients[variable])}{stars}"


def se_text(row: dict[str, str], variable: str) -> str:
    standard_errors = row.get("standard_errors", {})
    if not isinstance(standard_errors, dict) or variable not in standard_errors:
        return ""
    formatted = format_decimal(standard_errors[variable])
    return f"({formatted})" if formatted else ""


def ordered_variables(rows: list[dict[str, str]]) -> list[str]:
    out: list[str] = []
    for row in rows:
        for variable in [row["target"], *row["controls"].split()]:
            if variable and variable not in out:
                out.append(variable)
    return out


def set_run_font(run, *, bold: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def set_cell_text(cell, text: str, *, bold: bool = False, align: int | None = None, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_note_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge, value in (("top", "single"), ("left", "nil"), ("bottom", "nil"), ("right", "nil")):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), value)
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tc_pr.append(borders)


def add_table_title(doc: Document, table_num: int, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(f"Table {table_num}: {title}"))


def add_regression_table(doc: Document, table_num: int, title: str, rows: list[dict[str, str]]) -> None:
    add_table_title(doc, table_num, title)
    table = doc.add_table(rows=1, cols=len(rows) + 1)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, row in enumerate(rows, start=1):
        set_cell_text(table.rows[0].cells[idx], model_label(row, idx), align=WD_ALIGN_PARAGRAPH.CENTER)

    dep_row = table.add_row().cells
    set_cell_text(dep_row[0], "")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(dep_row[idx], row["depvar"], align=WD_ALIGN_PARAGRAPH.CENTER)

    for variable in ordered_variables(rows):
        coef_row = table.add_row().cells
        set_cell_text(coef_row[0], variable)
        for idx, row in enumerate(rows, start=1):
            set_cell_text(coef_row[idx], coef_text(row, variable), align=WD_ALIGN_PARAGRAPH.CENTER)
        se_row = table.add_row().cells
        set_cell_text(se_row[0], "")
        for idx, row in enumerate(rows, start=1):
            set_cell_text(se_row[idx], se_text(row, variable), align=WD_ALIGN_PARAGRAPH.CENTER)

    entity_fe_row = table.add_row().cells
    set_cell_text(entity_fe_row[0], "Entity FE")
    for idx in range(1, len(rows) + 1):
        set_cell_text(entity_fe_row[idx], "Yes", align=WD_ALIGN_PARAGRAPH.CENTER)

    time_fe_row = table.add_row().cells
    set_cell_text(time_fe_row[0], "Time FE")
    for idx in range(1, len(rows) + 1):
        set_cell_text(time_fe_row[idx], "Yes", align=WD_ALIGN_PARAGRAPH.CENTER)

    n_row = table.add_row().cells
    set_cell_text(n_row[0], "N")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(n_row[idx], format_n(row["nobs"]), align=WD_ALIGN_PARAGRAPH.CENTER)

    r2_row = table.add_row().cells
    set_cell_text(r2_row[0], "Adj. R2")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(r2_row[idx], format_decimal(row["r2"]), align=WD_ALIGN_PARAGRAPH.CENTER)

    note = "Standard errors in parentheses; *** p<0.01, ** p<0.05, * p<0.1"
    if rows and rows[0]["section"] == "iv_stage1" and rows[0].get("instrument"):
        note += f"; x ~ iv + cv; IV: {rows[0]['instrument']}"
    if rows and rows[0]["section"] == "iv_stage2" and rows[0].get("instrument"):
        note += f"; y ~ x_hat + cv; IV: {rows[0]['instrument']}"
    note_cells = table.add_row().cells
    note_cell = note_cells[0]
    if len(note_cells) > 1:
        note_cell = note_cell.merge(note_cells[-1])
    set_cell_text(note_cell, note, size=10)
    set_note_cell_borders(note_cell)


DESCRIPTIVE_HEADERS = ["VarName", "Obs", "Mean", "SD", "Min", "Median", "Max"]


def add_descriptive_table(doc: Document, table_num: int, stats_rows: list[list[str]]) -> None:
    add_table_title(doc, table_num, "描述性统计")
    table = doc.add_table(rows=1, cols=len(DESCRIPTIVE_HEADERS))
    table.style = "Table Grid"
    for idx, header in enumerate(DESCRIPTIVE_HEADERS):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for stats in stats_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], stats[0])
        for idx, value in enumerate(stats[1:], start=1):
            set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER)


def write_docx(
    path: Path,
    rows: list[dict[str, str]],
    metadata: dict[str, str],
    descriptive_rows: list[list[str]] | None = None,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    grouped: dict[str, list[dict[str, str]]] = {}
    for row in rows:
        key = doc_section_key(row["section"])
        grouped.setdefault(key, []).append(row)
    table_num = 0
    for table_num, (section_name, section_rows) in enumerate(grouped.items(), start=1):
        add_regression_table(doc, table_num, section_title(section_name), section_rows)
    if descriptive_rows:
        add_descriptive_table(doc, table_num + 1, descriptive_rows)

    doc.save(path)


def build_descriptive_var_pool(args: RegressionArgs, cv_subset: list[str]) -> list[str]:
    """Mirror the Stata env's desc_vars pool: raw analysis variables, deduplicated."""
    rob = parse_rob_vars(args.rob_vars)
    discrete = parse_discrete_values(args.het_disc_vals)
    active_het = [g for g in split_words(args.het_disc) if discrete.get(g)]
    items = [
        *split_words(args.y),
        *split_words(args.x),
        *cv_subset,
        *split_words(args.meds),
        *split_words(args.mods),
        *active_het,
        *split_words(args.iv),
        *split_words(rob.get("alt_x", "")),
        *split_words(rob.get("alt_y", "")),
    ]
    out: list[str] = []
    for item in items:
        if item and item not in out:
            out.append(item)
    return out


def build_descriptive_rows(df: pd.DataFrame, base_sample: pd.Series, variables: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for variable in variables:
        if variable not in df.columns:
            continue
        values = pd.to_numeric(df.loc[base_sample, variable], errors="coerce").dropna()
        if values.empty:
            continue
        rows.append(
            [
                variable,
                str(int(values.count())),
                format_decimal(values.mean()),
                format_decimal(values.std()),
                format_decimal(values.min()),
                format_decimal(values.median()),
                format_decimal(values.max()),
            ]
        )
    return rows


def run_final(
    values: dict[str, str],
    *,
    cv_idx: int,
    vce_idx: int,
    output_dir: str | None = None,
    source_path: str | None = None,
) -> dict[str, str]:
    args = RegressionArgs.from_mapping(values)
    result_dir = Path(output_dir) if output_dir else Path(os.environ.get("STARLANE_EXPORT", ".starlane"))
    result_dir.mkdir(parents=True, exist_ok=True)

    y_vars = split_words(args.y)
    x_vars = split_words(args.x)
    cv_all = split_words(args.cv)
    coef_direction = (args.coef_direction.strip().lower() or "positive")
    plan = build_model_plan(args)
    cv_subset = list(plan.cv_subset(cv_idx).controls)
    vce_choice = plan.vce_choice(vce_idx)

    optional_originals = split_words(args.meds) + split_words(args.mods) + split_words(args.iv)
    df = read_data(args.data_path)
    ensure_columns(df, [*y_vars, *x_vars, *cv_all, args.panelvar, args.timevar, *optional_originals])
    df = prepare_regression_data(df, args)
    df, panelvar = encode_panel_if_needed(df, args.panelvar)
    timevar = args.timevar
    specs = list(plan.specs_for_cv_idx(args, cv_idx))
    ensure_columns(df, [panelvar, timevar, *spec_required_columns(specs)])
    base_sample = make_base_sample(df, [panelvar, timevar, *spec_required_columns(specs)])

    rows: list[dict[str, str]] = []
    for spec in specs:
        spec_sample = apply_spec_condition(df, base_sample, spec)
        result = run_spec(df, spec, panelvar, timevar, vce_idx, spec_sample)
        stars = stars_for_result(result, coef_direction)
        rows.append(
            {
                "column": spec.column,
                "section": spec.section,
                "depvar": spec.depvar,
                "target": spec.target_var,
                "controls": " ".join(spec.controls),
                "condition": f"{spec.condition_var}={spec.condition_value}" if spec.condition_var else "",
                "coef": "" if result is None else f"{result.coef:.10g}",
                "se": "" if result is None else f"{result.se:.10g}",
                "p_value": "" if result is None else f"{result.p_value:.10g}",
                "stars": "*" * stars,
                "coef_with_stars": format_coef(result, stars),
                "nobs": "" if result is None else str(result.nobs),
                "r2": "" if result is None or not math.isfinite(result.r2) else f"{result.r2:.10g}",
                "coefficients": {} if result is None else result.coefficients,
                "standard_errors": {} if result is None else result.standard_errors,
                "p_values": {} if result is None else result.p_values,
                "instrument": spec.instrument,
            }
        )

    csv_path = result_dir / "final_result.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        fieldnames = ["column", "section", "depvar", "target", "controls", "condition", "coef", "se", "p_value", "stars", "coef_with_stars", "nobs", "r2"]
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows({field: row[field] for field in fieldnames} for row in rows)

    md_path = result_dir / "final_result.md"
    metadata = {
        "input": args.data_path,
        "cv_idx": str(cv_idx),
        "vce_idx": str(vce_idx),
        "vce_suffix": vce_choice.suffix,
        "cv_selected": " ".join(cv_subset),
    }
    md_lines = [
        "# Starlane Python Final Result",
        "",
        *[f"- {key}: `{value}`" for key, value in metadata.items()],
        "",
        "| section | depvar | target | controls | condition | coef | se | p | nobs | r2 |",
        "| --- | --- | --- | --- | --- | ---: | ---: | ---: | ---: | ---: |",
    ]
    for row in rows:
        md_lines.append(
            f"| {row['section']} | {row['depvar']} | {row['target']} | {row['controls']} | {row['condition']} | {row['coef_with_stars']} | {row['se']} | {row['p_value']} | {row['nobs']} | {row['r2']} |"
        )
    md_path.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    docx_path = result_dir / "final_result.docx"
    desc_vars = build_descriptive_var_pool(args, cv_subset)
    descriptive_rows = build_descriptive_rows(df, base_sample, desc_vars)
    write_docx(docx_path, rows, metadata, descriptive_rows)

    source_copy = copy_source_to_dir(Path(source_path or __file__), result_dir)
    return {
        "csv": str(csv_path),
        "markdown": str(md_path),
        "docx": str(docx_path),
        "source": str(source_copy),
    }
