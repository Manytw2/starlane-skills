"""Python final-stage env for a selected Starlane regression row."""

from __future__ import annotations

import csv
import json
import math
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from common import (
    RegressionArgs,
    apply_spec_condition,
    build_specs,
    compute_cv_subset,
    copy_source_to_dir,
    encode_panel_if_needed,
    ensure_columns,
    fail,
    format_coef,
    load_regression_args_json,
    load_selection_json,
    make_base_sample,
    prepare_regression_data,
    reject_positional_args,
    read_data,
    run_spec,
    spec_required_columns,
    split_words,
    stars_for_result,
    vce_suffix,
)


SECTION_TITLES = {
    "baseline": "基准回归",
    "robustness_alt_x": "稳健性检验-替换X",
    "robustness_alt_y": "稳健性检验-替换变量",
    "robustness_ln_x": "稳健性检验-X取对数",
    "robustness_ln_y": "稳健性检验-Y取对数",
    "robustness_lag": "稳健性检验-滞后期",
    "robustness_year": "稳健性检验-时间窗口",
    "iv_stage1": "工具变量-一阶段",
    "iv_stage2": "工具变量-二阶段 2SLS",
    "mediation": "中介机制",
    "moderation": "异质性分析-调节效应检验",
}


def section_title(section: str) -> str:
    if section in SECTION_TITLES:
        return SECTION_TITLES[section]
    if section.startswith("mediation_"):
        return "中介机制"
    if section.startswith("moderation_"):
        return f"Moderation: {section.removeprefix('moderation_')}"
    if section.startswith("heterogeneity_discrete_"):
        return f"Heterogeneity: {section.removeprefix('heterogeneity_discrete_')}"
    return section.replace("_", " ").title()


def doc_section_key(section: str) -> str:
    if section in ("baseline_nocv", "baseline_cv"):
        return "baseline"
    if section.startswith("mediation_"):
        return "mediation"
    if section.startswith("moderation_"):
        return "moderation"
    return section


def model_label(row: dict[str, str], idx: int) -> str:
    if row["condition"]:
        return row["condition"]
    return f"({idx})"


def format_decimal(value: str | float, digits: int = 3) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return ""
    if not math.isfinite(number):
        return ""
    return f"{number:.{digits}f}"


def format_n(value: str | int) -> str:
    try:
        number = int(float(value))
    except (TypeError, ValueError):
        return ""
    return f"{number:,}"


def coef_text(row: dict[str, str], variable: str) -> str:
    coefficients = row.get("coefficients", {})
    if not isinstance(coefficients, dict) or variable not in coefficients:
        return ""
    stars = row["stars"] if variable == row["target"] else ""
    return f"{format_decimal(coefficients[variable])}{stars}"


def se_text(row: dict[str, str], variable: str) -> str:
    standard_errors = row.get("standard_errors", {})
    if not isinstance(standard_errors, dict) or variable not in standard_errors:
        return ""
    formatted = format_decimal(standard_errors[variable])
    return f"({formatted})" if formatted else ""


def ordered_variables(rows: list[dict[str, str]]) -> list[str]:
    out: list[str] = []
    for row in rows:
        for variable in [row["target"], *row["controls"].split()]:
            if variable and variable not in out:
                out.append(variable)
    return out


def set_run_font(run, *, bold: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def set_cell_text(cell, text: str, *, bold: bool = False, align: int | None = None, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_note_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge, value in (("top", "single"), ("left", "nil"), ("bottom", "nil"), ("right", "nil")):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), value)
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tc_pr.append(borders)


def add_table_title(doc: Document, table_num: int, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(f"Table {table_num}: {title}"))


def add_regression_table(doc: Document, table_num: int, title: str, rows: list[dict[str, str]]) -> None:
    add_table_title(doc, table_num, title)
    table = doc.add_table(rows=1, cols=len(rows) + 1)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, row in enumerate(rows, start=1):
        set_cell_text(table.rows[0].cells[idx], model_label(row, idx), align=WD_ALIGN_PARAGRAPH.CENTER)

    dep_row = table.add_row().cells
    set_cell_text(dep_row[0], "")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(dep_row[idx], row["depvar"], align=WD_ALIGN_PARAGRAPH.CENTER)

    for variable in ordered_variables(rows):
        coef_row = table.add_row().cells
        set_cell_text(coef_row[0], variable)
        for idx, row in enumerate(rows, start=1):
            set_cell_text(coef_row[idx], coef_text(row, variable), align=WD_ALIGN_PARAGRAPH.CENTER)
        se_row = table.add_row().cells
        set_cell_text(se_row[0], "")
        for idx, row in enumerate(rows, start=1):
            set_cell_text(se_row[idx], se_text(row, variable), align=WD_ALIGN_PARAGRAPH.CENTER)

    entity_fe_row = table.add_row().cells
    set_cell_text(entity_fe_row[0], "Entity FE")
    for idx in range(1, len(rows) + 1):
        set_cell_text(entity_fe_row[idx], "Yes", align=WD_ALIGN_PARAGRAPH.CENTER)

    time_fe_row = table.add_row().cells
    set_cell_text(time_fe_row[0], "Time FE")
    for idx in range(1, len(rows) + 1):
        set_cell_text(time_fe_row[idx], "Yes", align=WD_ALIGN_PARAGRAPH.CENTER)

    n_row = table.add_row().cells
    set_cell_text(n_row[0], "N")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(n_row[idx], format_n(row["nobs"]), align=WD_ALIGN_PARAGRAPH.CENTER)

    r2_row = table.add_row().cells
    set_cell_text(r2_row[0], "Adj. R2")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(r2_row[idx], format_decimal(row["r2"]), align=WD_ALIGN_PARAGRAPH.CENTER)

    note = "Standard errors in parentheses; *** p<0.01, ** p<0.05, * p<0.1"
    if rows and rows[0]["section"] == "iv_stage1" and rows[0].get("instrument"):
        note += f"; x ~ iv + cv; IV: {rows[0]['instrument']}"
    if rows and rows[0]["section"] == "iv_stage2" and rows[0].get("instrument"):
        note += f"; y ~ x_hat + cv; IV: {rows[0]['instrument']}"
    note_cells = table.add_row().cells
    note_cell = note_cells[0]
    if len(note_cells) > 1:
        note_cell = note_cell.merge(note_cells[-1])
    set_cell_text(note_cell, note, size=10)
    set_note_cell_borders(note_cell)


def write_docx(path: Path, rows: list[dict[str, str]], metadata: dict[str, str]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    grouped: dict[str, list[dict[str, str]]] = {}
    for row in rows:
        key = doc_section_key(row["section"])
        grouped.setdefault(key, []).append(row)
    for table_num, (section_name, section_rows) in enumerate(grouped.items(), start=1):
        add_regression_table(doc, table_num, section_title(section_name), section_rows)

    doc.save(path)


def run_final(
    values: dict[str, str],
    *,
    cv_idx: int,
    vce_idx: int,
    output_arg: str | None = None,
    source_path: str | None = None,
) -> dict[str, str]:
    if len(values) < 18:
        raise ValueError("Expected 18 summary args")
    args = RegressionArgs.from_mapping(values)
    result_dir = Path(output_arg) if output_arg else Path(os.environ.get("STARLANE_EXPORT", ".starlane"))
    result_dir.mkdir(parents=True, exist_ok=True)

    y_vars = split_words(args.y)
    x_vars = split_words(args.x)
    cv_all = split_words(args.cv)
    cv_fixed = split_words(args.cv_fixed)
    cv_min_count = int(args.cv_min_count.strip() or "0")
    coef_direction = (args.coef_direction.strip().lower() or "positive")
    cv_subset = compute_cv_subset(cv_all, cv_fixed, cv_min_count, cv_idx)

    optional_originals = split_words(args.meds) + split_words(args.mods) + split_words(args.iv)
    df = read_data(args.input_dta)
    ensure_columns(df, [*y_vars, *x_vars, *cv_all, args.panelvar, args.timevar, *optional_originals])
    df = prepare_regression_data(df, args)
    df, panelvar = encode_panel_if_needed(df, args.panelvar)
    timevar = args.timevar
    specs = build_specs(args, cv_subset)
    ensure_columns(df, [panelvar, timevar, *spec_required_columns(specs)])
    base_sample = make_base_sample(df, [panelvar, timevar, *spec_required_columns(specs)])

    rows: list[dict[str, str]] = []
    for spec in specs:
        spec_sample = apply_spec_condition(df, base_sample, spec)
        result = run_spec(df, spec, panelvar, timevar, vce_idx, spec_sample)
        stars = stars_for_result(result, coef_direction)
        rows.append(
            {
                "column": spec.column,
                "section": spec.section,
                "depvar": spec.depvar,
                "target": spec.target_var,
                "controls": " ".join(spec.controls),
                "condition": f"{spec.condition_var}={spec.condition_value}" if spec.condition_var else "",
                "coef": "" if result is None else f"{result.coef:.10g}",
                "se": "" if result is None else f"{result.se:.10g}",
                "p_value": "" if result is None else f"{result.p_value:.10g}",
                "stars": "*" * stars,
                "coef_with_stars": format_coef(result, stars),
                "nobs": "" if result is None else str(result.nobs),
                "r2": "" if result is None else f"{result.r2:.10g}",
                "coefficients": {} if result is None else result.coefficients,
                "standard_errors": {} if result is None else result.standard_errors,
                "instrument": spec.instrument,
            }
        )

    csv_path = result_dir / "final_result.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        fieldnames = ["column", "section", "depvar", "target", "controls", "condition", "coef", "se", "p_value", "stars", "coef_with_stars", "nobs", "r2"]
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows({field: row[field] for field in fieldnames} for row in rows)

    md_path = result_dir / "final_result.md"
    metadata = {
        "input": args.input_dta,
        "cv_idx": str(cv_idx),
        "vce_idx": str(vce_idx),
        "vce_suffix": vce_suffix(vce_idx, panelvar, timevar),
        "cv_selected": " ".join(cv_subset),
    }
    md_lines = [
        "# Starlane Python Final Result",
        "",
        *[f"- {key}: `{value}`" for key, value in metadata.items()],
        "",
        "| section | depvar | target | controls | condition | coef | se | p | nobs | r2 |",
        "| --- | --- | --- | --- | --- | ---: | ---: | ---: | ---: | ---: |",
    ]
    for row in rows:
        md_lines.append(
            f"| {row['section']} | {row['depvar']} | {row['target']} | {row['controls']} | {row['condition']} | {row['coef_with_stars']} | {row['se']} | {row['p_value']} | {row['nobs']} | {row['r2']} |"
        )
    md_path.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    docx_path = result_dir / "final_result.docx"
    write_docx(docx_path, rows, metadata)

    run_note = result_dir / "python_env_run_note.md"
    run_note.write_text(
        "\n".join(
            [
                "# Python Env Run Note",
                "",
                "This Python env implements the Starlane summary/final workflow contract across the same section families as the Stata env.",
                "",
                "It uses an internal numpy OLS implementation with absorbed panel and time fixed effects.",
                "Numerical equivalence with Stata reghdfe is not guaranteed until cross-env tests are added.",
                "",
                "Raw args:",
                "",
                "```json",
                json.dumps({"args": args.as_mapping(), "selection": {"cv_idx": cv_idx, "vce_idx": vce_idx}, "result_dir": str(result_dir)}, ensure_ascii=False, indent=2),
                "```",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    source_copy = copy_source_to_dir(Path(source_path or __file__), result_dir)
    return {
        "csv": str(csv_path),
        "markdown": str(md_path),
        "docx": str(docx_path),
        "source": str(source_copy),
        "run_note": str(run_note),
    }


def main() -> int:
    try:
        reject_positional_args(sys.argv)
        import argparse

        parser = argparse.ArgumentParser(description="Run Starlane Python final stage from JSON files.")
        parser.add_argument("--args-json", required=True, help="Path to regression_args.json")
        parser.add_argument("--selection-json", required=True, help="Path to selected_candidate.json with cv_idx and vce_idx")
        parser.add_argument("--output", help="Output directory. Defaults to STARLANE_EXPORT or .starlane")
        ns = parser.parse_args(sys.argv[1:])

        values = load_regression_args_json(ns.args_json)
        selection = load_selection_json(ns.selection_json)
        outputs = run_final(values, cv_idx=selection["cv_idx"], vce_idx=selection["vce_idx"], output_arg=ns.output)
        print(f"STARLANE_FINAL_OUTPUT: {outputs['docx']}")
        print(f"STARLANE_SOURCE_ARTIFACT: {outputs['source']}")
        return 0
    except Exception as e:
        return fail(str(e))


if __name__ == "__main__":
    sys.exit(main())
