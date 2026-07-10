"""Final stage (python env): regression args + selection -> 可复现源码.

IN:  regression_args.json + selected_candidate.json（cv_idx / vce_idx）
OUT: 自包含的 Python 源码（只依赖 pandas/numpy/pyfixest/python-docx 公开包，
     运行后产出 final_result.* 到 STARLANE_EXPORT 或脚本所在目录）

The generated file expands every regression spec at generation time (mirroring
the Stata env's template expansion) so users can read, run, and adapt it
without this repository.
"""

from __future__ import annotations

import sys
from pathlib import Path

WORKFLOW_SCRIPTS = Path(__file__).resolve().parents[2] / "workflow"
if str(WORKFLOW_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(WORKFLOW_SCRIPTS))

from contracts import REGRESSION_ARG_NAMES, load_regression_args_json, load_selection_json  # noqa: E402
from model_plan import (  # noqa: E402
    RegressionArgsProxy,
    build_model_plan,
    parse_bool_default_yes,
    parse_discrete_values,
    parse_rob_vars,
    sample_pool_columns,
    split_words,
)


def parse_args(argv: list[str]) -> tuple[dict[str, str], int, int, Path]:
    import argparse

    parser = argparse.ArgumentParser(description="Generate a self-contained Starlane Python final-stage source file.")
    parser.add_argument("--args-json", required=True, help="Path to regression_args.json")
    parser.add_argument("--selection-json", required=True, help="Path to selected_candidate.json with cv_idx and vce_idx")
    parser.add_argument("--output", required=True, help="Output generated Python source path")
    ns = parser.parse_args(argv[1:])
    mapping = load_regression_args_json(ns.args_json)
    selection = load_selection_json(ns.selection_json)
    values = {name: str(mapping[name]) for name in REGRESSION_ARG_NAMES}
    return values, int(selection["cv_idx"]), int(selection["vce_idx"]), Path(ns.output)


def unique_preserve(items: list[str]) -> list[str]:
    out: list[str] = []
    for item in items:
        if item and item not in out:
            out.append(item)
    return out


def build_desc_vars(values: dict[str, str], cv_subset: list[str]) -> list[str]:
    """Descriptive-stats variable pool (raw analysis variables, deduplicated)."""
    rob = parse_rob_vars(values["rob_vars"])
    discrete = parse_discrete_values(values["het_disc_vals"])
    active_het = [g for g in split_words(values["het_disc"]) if discrete.get(g)]
    return unique_preserve(
        [
            *split_words(values["y"]),
            *split_words(values["x"]),
            *cv_subset,
            *split_words(values["meds"]),
            *split_words(values["mods"]),
            *active_het,
            *split_words(values["iv"]),
            *split_words(rob.get("alt_x", "")),
            *split_words(rob.get("alt_y", "")),
        ]
    )


def build_ln_sources(values: dict[str, str], key: str, base_vars: list[str]) -> list[str]:
    rob = parse_rob_vars(values["rob_vars"])
    sources: list[str] = []
    if parse_bool_default_yes(values[key]):
        sources.extend(base_vars)
    sources.extend(split_words(rob.get(key, "")))
    return unique_preserve(sources)


def build_lag_periods(values: dict[str, str]) -> list[int]:
    rob = parse_rob_vars(values["rob_vars"])
    periods: list[int] = []
    for period in split_words(rob.get("lag", "")):
        try:
            periods.append(int(period))
        except ValueError:
            continue
    return periods


def render_specs_literal(specs) -> str:
    """Render the spec list as a readable Python literal grouped by section."""
    lines: list[str] = ["SPECS = ["]
    last_section = None
    for spec in specs:
        if spec.section != last_section:
            lines.append(f"    # section: {spec.section}")
            last_section = spec.section
        parts = [
            f"column={spec.column!r}",
            f"section={spec.section!r}",
            f"depvar={spec.depvar!r}",
            f"target={spec.target_var!r}",
            f"controls={list(spec.controls)!r}",
        ]
        if spec.condition_var:
            parts.append(f"condition_var={spec.condition_var!r}")
            parts.append(f"condition_value={spec.condition_value!r}")
        if spec.instrument:
            parts.append(f"instrument={spec.instrument!r}")
        lines.append(f"    dict({', '.join(parts)}),")
    lines.append("]")
    return "\n".join(lines)


def render_list(name: str, items: list) -> str:
    return f"{name} = {items!r}"


TEMPLATE = '''"""Generated Starlane Python regression final-stage source.

Self-contained reproducibility script. Dependencies (all public packages):

    pip install pandas numpy pyfixest python-docx

Running this file re-estimates every regression below with pyfixest
(high-dimensional fixed effects, aligned with Stata reghdfe) and writes
final_result.csv / final_result.md / final_result.docx to the directory in
the STARLANE_EXPORT environment variable, or next to this file by default.
"""

from __future__ import annotations

import csv
import math
import os
import shutil
from pathlib import Path

import numpy as np
import pandas as pd
import pyfixest as pf

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# ---------------------------------------------------------------------------
# Parameters (expanded from the confirmed analysis plan and selected candidate)
# ---------------------------------------------------------------------------

__PARAMS__

# Every regression run by this script, one entry per model column.
# depvar ~ target + controls, absorbing entity and time fixed effects.
__SPECS__


# ---------------------------------------------------------------------------
# Data loading and preparation
# ---------------------------------------------------------------------------

def read_data(path: str) -> pd.DataFrame:
    suffix = Path(path).suffix.lower()
    if suffix == ".dta":
        return pd.read_stata(path)
    if suffix == ".csv":
        return pd.read_csv(path)
    if suffix in (".xls", ".xlsx"):
        return pd.read_excel(path)
    raise ValueError(f"Unsupported input format: {path}")


def ensure_columns(df: pd.DataFrame, columns: list[str]) -> None:
    missing = [c for c in columns if c and c not in df.columns]
    if missing:
        raise ValueError(f"Missing variable(s): {', '.join(missing)}")


def unique_preserve(items: list[str]) -> list[str]:
    out: list[str] = []
    for item in items:
        if item and item not in out:
            out.append(item)
    return out


def add_ln_column(df: pd.DataFrame, source: str) -> None:
    name = f"ln_{source}"
    if name not in df.columns and source in df.columns:
        values = pd.to_numeric(df[source], errors="coerce")
        df[name] = np.where(values > 0, np.log(values), np.nan)


def add_std_column(df: pd.DataFrame, source: str) -> None:
    values = pd.to_numeric(df[source], errors="coerce")
    std = values.std()
    df[f"std_{source}"] = (values - values.mean()) / std if std and np.isfinite(std) else np.nan


def prepare_data(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for source in LN_X_SOURCES:
        add_ln_column(out, source)
    for source in LN_Y_SOURCES:
        add_ln_column(out, source)
    if LAG_PERIODS:
        out = out.sort_values([PANELVAR, TIMEVAR]).copy()
        for period in LAG_PERIODS:
            for x in X_VARS:
                out[f"l{period}_{x}"] = out.groupby(PANELVAR, sort=False)[x].shift(period)
    for x in X_VARS:
        add_std_column(out, x)
    for mod in MOD_VARS:
        if mod in out.columns:
            add_std_column(out, mod)
            for x in X_VARS:
                out[f"inter_{x}_{mod}"] = out[f"std_{x}"] * out[f"std_{mod}"]
    return out


def encode_panel_if_needed(df: pd.DataFrame, panelvar: str) -> tuple[pd.DataFrame, str]:
    """Numeric-encode a string panel id (mirrors Stata egen group())."""
    if pd.api.types.is_numeric_dtype(df[panelvar]):
        return df, panelvar
    out = df.copy()
    new_name = f"{panelvar}_gid"
    suffix = 0
    while new_name in out.columns:
        suffix += 1
        new_name = f"{panelvar}_g{suffix}"
    out[new_name] = pd.factorize(out[panelvar], sort=True)[0]
    return out, new_name


def apply_condition(df: pd.DataFrame, sample: pd.Series, condition_var: str, raw: str) -> pd.Series:
    series = df[condition_var]
    if ":" in raw:
        left, right = [v.strip() for v in raw.split(":", 1)]
        try:
            lo, hi = float(left), float(right)
            numeric = pd.to_numeric(series, errors="coerce")
            return sample & (numeric >= lo) & (numeric <= hi)
        except ValueError:
            pass
    try:
        value = float(raw)
        return sample & (pd.to_numeric(series, errors="coerce") == value)
    except ValueError:
        return sample & (series.astype(str) == raw)


# ---------------------------------------------------------------------------
# Estimation (pyfixest.feols with entity + time fixed effects)
# ---------------------------------------------------------------------------

def vcov_choice(vce_idx: int, panelvar: str, timevar: str):
    if vce_idx == 0:
        return "iid"
    if vce_idx == 1:
        return "hetero"
    if vce_idx == 2:
        return {"CRV1": panelvar}
    if vce_idx == 3:
        return {"CRV1": f"{panelvar} + {timevar}"}
    raise ValueError("vce_idx must be 0-3")


def quote_name(name: str) -> str:
    return "`" + name.replace("`", "\\\\`") + "`"


def formula_terms(names: list[str]) -> str:
    return " + ".join(quote_name(n) for n in names) if names else "1"


def model_frame(df: pd.DataFrame, sample: pd.Series, columns: list[str], numeric_columns: list[str]) -> pd.DataFrame:
    data = df.loc[sample, columns].copy().reset_index(drop=True)
    for col in numeric_columns:
        data[col] = pd.to_numeric(data[col], errors="coerce")
    return data.replace([np.inf, -np.inf], np.nan).dropna()


def series_value(values: pd.Series, name: str) -> float:
    if name in values.index:
        return float(values.loc[name])
    for key, value in values.items():
        if str(key) == str(name):
            return float(value)
    return math.nan


def fit_spec(df: pd.DataFrame, spec: dict, panelvar: str, sample: pd.Series) -> dict | None:
    """Fit one regression spec; return None when the model is not estimable."""
    controls = list(spec["controls"])
    if spec.get("instrument"):
        # IV second stage: depvar ~ controls | FE | target ~ instrument
        columns = [spec["depvar"], spec["target"], spec["instrument"], *controls, panelvar, TIMEVAR]
        numeric = [spec["depvar"], spec["target"], spec["instrument"], *controls]
        formula = (
            f"{quote_name(spec['depvar'])} ~ {formula_terms(controls)}"
            f" | {quote_name(panelvar)} + {quote_name(TIMEVAR)}"
            f" | {quote_name(spec['target'])} ~ {quote_name(spec['instrument'])}"
        )
        variance_check = [spec["target"], spec["instrument"], *controls]
    else:
        regressors = [spec["target"], *controls]
        columns = [spec["depvar"], *regressors, panelvar, TIMEVAR]
        numeric = [spec["depvar"], *regressors]
        formula = f"{quote_name(spec['depvar'])} ~ {formula_terms(regressors)} | {quote_name(panelvar)} + {quote_name(TIMEVAR)}"
        variance_check = regressors

    data = model_frame(df, sample, columns, numeric)
    if len(data) < 30:
        return None
    if any(data[col].std() == 0 for col in variance_check):
        return None
    try:
        model = pf.feols(fml=formula, data=data, vcov=vcov_choice(VCE_IDX, panelvar, TIMEVAR))
    except Exception:
        return None

    coefs = model.coef()
    ses = model.se()
    p_values = model.pvalue()
    coef = series_value(coefs, spec["target"])
    if not math.isfinite(coef):
        return None
    try:
        nobs = int(getattr(model, "_N"))
    except (AttributeError, TypeError, ValueError):
        nobs = len(data)
    try:
        r2 = float(getattr(model, "_adj_r2", math.nan))
    except (TypeError, ValueError):
        r2 = math.nan
    return {
        "coef": coef,
        "se": series_value(ses, spec["target"]),
        "p_value": series_value(p_values, spec["target"]),
        "nobs": nobs,
        "r2": r2 if math.isfinite(r2) else math.nan,
        "coefficients": {str(k): float(v) for k, v in coefs.items() if math.isfinite(float(v))},
        "standard_errors": {str(k): float(v) for k, v in ses.items() if math.isfinite(float(v))},
        "p_values": {str(k): float(v) for k, v in p_values.items() if math.isfinite(float(v))},
    }


def stars_for(coef: float, p_value: float) -> int:
    direction_ok = coef > 0 if COEF_DIRECTION == "positive" else coef < 0
    if not direction_ok or not math.isfinite(p_value):
        return 0
    if p_value < 0.01:
        return 3
    if p_value < 0.05:
        return 2
    if p_value < 0.1:
        return 1
    return 0


# ---------------------------------------------------------------------------
# Word table rendering (booktabs-style regression tables)
# ---------------------------------------------------------------------------

SECTION_TITLES = {
    "baseline": "基准回归",
    "rob_alt_x": "稳健性检验-替换X",
    "rob_alt_y": "稳健性检验-替换变量",
    "rob_ln_x": "稳健性检验-X取对数",
    "rob_ln_y": "稳健性检验-Y取对数",
    "rob_lag": "稳健性检验-滞后期",
    "rob_year": "稳健性检验-时间窗口",
    "iv_stage1": "工具变量-一阶段",
    "iv_stage2": "工具变量-二阶段 2SLS",
    "med": "中介机制",
    "mod": "异质性分析-调节效应检验",
}


def section_title(section: str) -> str:
    if section in SECTION_TITLES:
        return SECTION_TITLES[section]
    if section.startswith("het_disc_"):
        return f"Heterogeneity: {section.removeprefix('het_disc_')}"
    return section.replace("_", " ").title()


def doc_section_key(section: str) -> str:
    if section in ("baseline_nocv", "baseline_cv"):
        return "baseline"
    if section.startswith("med_"):
        return "med"
    if section.startswith("mod_"):
        return "mod"
    return section


def model_label(row: dict, idx: int) -> str:
    return row["condition"] if row["condition"] else f"({idx})"


def format_decimal(value, digits: int = 3) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return ""
    if not math.isfinite(number):
        return ""
    return f"{number:.{digits}f}"


def format_n(value) -> str:
    try:
        number = int(float(value))
    except (TypeError, ValueError):
        return ""
    return f"{number:,}"


def stars_from_p(p_value) -> str:
    try:
        p = float(p_value)
    except (TypeError, ValueError):
        return ""
    if not math.isfinite(p):
        return ""
    if p < 0.01:
        return "***"
    if p < 0.05:
        return "**"
    if p < 0.1:
        return "*"
    return ""


def coef_text(row: dict, variable: str) -> str:
    coefficients = row.get("coefficients", {})
    if variable not in coefficients:
        return ""
    if variable == row["target"]:
        stars = row["stars"]
    else:
        stars = stars_from_p(row.get("p_values", {}).get(variable))
    return f"{format_decimal(coefficients[variable])}{stars}"


def se_text(row: dict, variable: str) -> str:
    standard_errors = row.get("standard_errors", {})
    if variable not in standard_errors:
        return ""
    formatted = format_decimal(standard_errors[variable])
    return f"({formatted})" if formatted else ""


def ordered_variables(rows: list[dict]) -> list[str]:
    out: list[str] = []
    for row in rows:
        for variable in [row["target"], *row["controls"].split()]:
            if variable and variable not in out:
                out.append(variable)
    return out


def set_run_font(run, *, bold: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def set_cell_text(cell, text: str, *, bold: bool = False, align=None, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_borders(cell, edges: dict) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        value, size = edges.get(edge, ("nil", "0"))
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), value)
        if value != "nil":
            element.set(qn("w:sz"), size)
            element.set(qn("w:color"), "000000")
        borders.append(element)
    tc_pr.append(borders)


# Booktabs-style rule weights in eighths of a point: heavy top/bottom, light midrule.
RULE_HEAVY = "12"
RULE_LIGHT = "6"


def set_row_rules(row, *, top: str | None = None, bottom: str | None = None) -> None:
    edges = {}
    if top:
        edges["top"] = ("single", top)
    if bottom:
        edges["bottom"] = ("single", bottom)
    for cell in row.cells:
        _set_cell_borders(cell, edges)


def add_table_title(doc: Document, table_num: int, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(f"Table {table_num}: {title}"))


def add_regression_table(doc: Document, table_num: int, title: str, rows: list[dict]) -> None:
    add_table_title(doc, table_num, title)
    table = doc.add_table(rows=1, cols=len(rows) + 1)
    set_cell_text(table.rows[0].cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, row in enumerate(rows, start=1):
        set_cell_text(table.rows[0].cells[idx], model_label(row, idx), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_row_rules(table.rows[0], top=RULE_HEAVY)

    dep_row_obj = table.add_row()
    dep_row = dep_row_obj.cells
    set_cell_text(dep_row[0], "")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(dep_row[idx], row["depvar"], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_row_rules(dep_row_obj, bottom=RULE_LIGHT)

    for variable in ordered_variables(rows):
        coef_row = table.add_row().cells
        set_cell_text(coef_row[0], variable)
        for idx, row in enumerate(rows, start=1):
            set_cell_text(coef_row[idx], coef_text(row, variable), align=WD_ALIGN_PARAGRAPH.CENTER)
        se_row = table.add_row().cells
        set_cell_text(se_row[0], "")
        for idx, row in enumerate(rows, start=1):
            set_cell_text(se_row[idx], se_text(row, variable), align=WD_ALIGN_PARAGRAPH.CENTER)

    for label in ("Entity FE", "Time FE"):
        fe_row = table.add_row().cells
        set_cell_text(fe_row[0], label)
        for idx in range(1, len(rows) + 1):
            set_cell_text(fe_row[idx], "Yes", align=WD_ALIGN_PARAGRAPH.CENTER)

    n_row = table.add_row().cells
    set_cell_text(n_row[0], "N")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(n_row[idx], format_n(row["nobs"]), align=WD_ALIGN_PARAGRAPH.CENTER)

    r2_row = table.add_row().cells
    set_cell_text(r2_row[0], "Adj. R2")
    for idx, row in enumerate(rows, start=1):
        set_cell_text(r2_row[idx], format_decimal(row["r2"]), align=WD_ALIGN_PARAGRAPH.CENTER)

    note = "Standard errors in parentheses; *** p<0.01, ** p<0.05, * p<0.1"
    if rows and rows[0]["section"] == "iv_stage1" and rows[0].get("instrument"):
        note += f"; x ~ iv + cv; IV: {rows[0]['instrument']}"
    if rows and rows[0]["section"] == "iv_stage2" and rows[0].get("instrument"):
        note += f"; y ~ x_hat + cv; IV: {rows[0]['instrument']}"
    note_cells = table.add_row().cells
    note_cell = note_cells[0]
    if len(note_cells) > 1:
        note_cell = note_cell.merge(note_cells[-1])
    set_cell_text(note_cell, note, size=10)
    _set_cell_borders(note_cell, {"top": ("single", RULE_HEAVY)})


DESCRIPTIVE_HEADERS = ["VarName", "Obs", "Mean", "SD", "Min", "Median", "Max"]


def add_descriptive_table(doc: Document, table_num: int, stats_rows: list[list[str]]) -> None:
    add_table_title(doc, table_num, "描述性统计")
    table = doc.add_table(rows=1, cols=len(DESCRIPTIVE_HEADERS))
    for idx, header in enumerate(DESCRIPTIVE_HEADERS):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_row_rules(table.rows[0], top=RULE_HEAVY, bottom=RULE_LIGHT)
    for stats in stats_rows:
        row_obj = table.add_row()
        cells = row_obj.cells
        set_cell_text(cells[0], stats[0])
        for idx, value in enumerate(stats[1:], start=1):
            set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_row_rules(table.rows[-1], bottom=RULE_HEAVY)


def write_docx(path: Path, rows: list[dict], descriptive_rows: list[list[str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    grouped: dict[str, list[dict]] = {}
    for row in rows:
        grouped.setdefault(doc_section_key(row["section"]), []).append(row)
    table_num = 0
    for table_num, (section_name, section_rows) in enumerate(grouped.items(), start=1):
        add_regression_table(doc, table_num, section_title(section_name), section_rows)
    if descriptive_rows:
        add_descriptive_table(doc, table_num + 1, descriptive_rows)
    doc.save(path)


def build_descriptive_rows(df: pd.DataFrame, base_sample: pd.Series) -> list[list[str]]:
    rows: list[list[str]] = []
    for variable in DESC_VARS:
        if variable not in df.columns:
            continue
        values = pd.to_numeric(df.loc[base_sample, variable], errors="coerce").dropna()
        if values.empty:
            continue
        rows.append(
            [
                variable,
                str(int(values.count())),
                format_decimal(values.mean()),
                format_decimal(values.std()),
                format_decimal(values.min()),
                format_decimal(values.median()),
                format_decimal(values.max()),
            ]
        )
    return rows


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    result_dir = Path(os.environ.get("STARLANE_EXPORT") or Path(__file__).resolve().parent)
    result_dir.mkdir(parents=True, exist_ok=True)

    df = read_data(DATA_PATH)
    ensure_columns(df, [PANELVAR, TIMEVAR, *SAMPLE_POOL_VARS])
    df = prepare_data(df)
    df, panelvar = encode_panel_if_needed(df, PANELVAR)

    required = unique_preserve(
        [c for spec in SPECS for c in [spec["depvar"], spec["target"], *spec["controls"], spec.get("condition_var", "")]]
    )
    ensure_columns(df, [panelvar, TIMEVAR, *required])

    # Shared base sample anchored on raw variables only; derived-column specs
    # (lag/ln/std) lose rows inside their own regression instead.
    base_sample = df[[panelvar, TIMEVAR, *SAMPLE_POOL_VARS]].notna().all(axis=1)

    rows: list[dict] = []
    for spec in SPECS:
        sample = base_sample
        if spec.get("condition_var"):
            sample = apply_condition(df, base_sample, spec["condition_var"], spec["condition_value"])
        result = fit_spec(df, spec, panelvar, sample)
        stars = 0 if result is None else stars_for(result["coef"], result["p_value"])
        rows.append(
            {
                "column": spec["column"],
                "section": spec["section"],
                "depvar": spec["depvar"],
                "target": spec["target"],
                "controls": " ".join(spec["controls"]),
                "condition": f"{spec['condition_var']}={spec['condition_value']}" if spec.get("condition_var") else "",
                "coef": "" if result is None else f"{result['coef']:.10g}",
                "se": "" if result is None else f"{result['se']:.10g}",
                "p_value": "" if result is None else f"{result['p_value']:.10g}",
                "stars": "*" * stars,
                "coef_with_stars": "" if result is None else f"{result['coef']:.6g}{'*' * stars}",
                "nobs": "" if result is None else str(result["nobs"]),
                "r2": "" if result is None or not math.isfinite(result["r2"]) else f"{result['r2']:.10g}",
                "coefficients": {} if result is None else result["coefficients"],
                "standard_errors": {} if result is None else result["standard_errors"],
                "p_values": {} if result is None else result["p_values"],
                "instrument": spec.get("instrument", ""),
            }
        )

    csv_path = result_dir / "final_result.csv"
    fieldnames = ["column", "section", "depvar", "target", "controls", "condition", "coef", "se", "p_value", "stars", "coef_with_stars", "nobs", "r2"]
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows({field: row[field] for field in fieldnames} for row in rows)

    md_path = result_dir / "final_result.md"
    metadata = {
        "input": DATA_PATH,
        "cv_idx": str(CV_IDX),
        "vce_idx": str(VCE_IDX),
        "vce_suffix": VCE_SUFFIX,
        "cv_selected": " ".join(CV_SELECTED),
    }
    md_lines = [
        "# Starlane Python Final Result",
        "",
        *[f"- {key}: `{value}`" for key, value in metadata.items()],
        "",
        "| section | depvar | target | controls | condition | coef | se | p | nobs | r2 |",
        "| --- | --- | --- | --- | --- | ---: | ---: | ---: | ---: | ---: |",
    ]
    for row in rows:
        md_lines.append(
            f"| {row['section']} | {row['depvar']} | {row['target']} | {row['controls']} | {row['condition']} | {row['coef_with_stars']} | {row['se']} | {row['p_value']} | {row['nobs']} | {row['r2']} |"
        )
    md_path.write_text("\\n".join(md_lines) + "\\n", encoding="utf-8")

    docx_path = result_dir / "final_result.docx"
    write_docx(docx_path, rows, build_descriptive_rows(df, base_sample))

    source = Path(__file__).resolve()
    source_copy = result_dir / source.name
    if source != source_copy.resolve():
        shutil.copy2(source, source_copy)

    print(f"STARLANE_FINAL_OUTPUT: {docx_path}")
    print(f"STARLANE_SOURCE_ARTIFACT: {source_copy}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
'''


def build_source(values: dict[str, str], cv_idx: int, vce_idx: int) -> str:
    args_proxy = RegressionArgsProxy(values)
    plan = build_model_plan(args_proxy)
    cv_subset = list(plan.cv_subset(cv_idx).controls)
    vce_choice = plan.vce_choice(vce_idx)
    specs = plan.specs_for_cv_idx(args_proxy, cv_idx)

    x_vars = split_words(values["x"])
    params = "\n".join(
        [
            f"DATA_PATH = {values['data_path']!r}",
            f"PANELVAR = {values['panelvar']!r}",
            f"TIMEVAR = {values['timevar']!r}",
            f"CV_IDX = {cv_idx}  # selected control-variable combination",
            f"VCE_IDX = {vce_idx}  # 0=iid, 1=hetero, 2=cluster(entity), 3=cluster(entity+time)",
            f"VCE_SUFFIX = {vce_choice.suffix!r}",
            f"COEF_DIRECTION = {values['coef_direction'].strip().lower() or 'positive'!r}",
            render_list("CV_SELECTED", cv_subset),
            render_list("X_VARS", x_vars),
            render_list("SAMPLE_POOL_VARS", sample_pool_columns(args_proxy, cv_subset)),
            render_list("DESC_VARS", build_desc_vars(values, cv_subset)),
            render_list("LN_X_SOURCES", build_ln_sources(values, "ln_x", x_vars)),
            render_list("LN_Y_SOURCES", build_ln_sources(values, "ln_y", split_words(values["y"]))),
            render_list("LAG_PERIODS", build_lag_periods(values)),
            render_list("MOD_VARS", split_words(values["mods"])),
        ]
    )

    specs_literal = render_specs_literal(specs)
    generated = TEMPLATE.replace("__PARAMS__", params).replace("__SPECS__", specs_literal)

    generated_spec_count = specs_literal.count("dict(")
    if generated_spec_count != len(specs):
        raise ValueError(f"Generated spec count {generated_spec_count} does not match ModelPlan spec count {len(specs)}")
    return generated


def main() -> int:
    try:
        values, cv_idx, vce_idx, output = parse_args(sys.argv)
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(build_source(values, cv_idx, vce_idx), encoding="utf-8")
        output.chmod(0o755)
        print(f"Generated: {output}")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
