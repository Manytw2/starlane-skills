"""Regression backend smoke and contract test for the demo dataset."""

from __future__ import annotations

import csv
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "skills" / "starlane-regression" / "scripts"
sys.path.insert(0, str(SCRIPTS))
from regression_backend_common import RegressionArgs, compute_cv_subsets, split_words

DEMO_DTA = Path("/Users/daydream/Desktop/demo.dta")
OUT = ROOT / ".starlane-test" / "python-demo"
PY_OUT = OUT / "python"
STATA_OUT = OUT / "stata"
UV_PYTHON = ["uv", "run", "python"]

SUMMARY_ARGS = [
    str(DEMO_DTA),
    "lnApplyG lnGrantG",
    "Attention",
    "Scale Lev lnAge Tange Cash ROA SOE Top1 Inst",
    "Scale Lev lnAge ROA",
    "5",
    "id",
    "year",
    "Charge|Subsidy|lnCSR",
    "OverSea|lnMediaPos|lnMediaNeg",
    "",
    "",
    "alt_y:lnAGreenInv lnGGreenInv|lag:1",
    "0",
    "0",
    "",
    "Thermalinv",
    "positive",
    "0",
    "1",
    "",
]


def summary_chunk_path(out_dir: Path) -> Path:
    start = SUMMARY_ARGS[18] if len(SUMMARY_ARGS) > 18 else ""
    end = SUMMARY_ARGS[19] if len(SUMMARY_ARGS) > 19 else ""
    if start != "" and end != "":
        return out_dir / "tmp" / f"combination_summary_part_{start}_{end}.csv"
    return out_dir / "combination_summary.csv"


def run(cmd: list[str], env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    result = subprocess.run(cmd, cwd=ROOT, env=env, text=True, capture_output=True)
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr, file=sys.stderr)
        raise SystemExit(result.returncode)
    if result.stderr.strip():
        print(result.stderr, file=sys.stderr)
        raise SystemExit("Expected clean stderr")
    return result


def find_stata() -> str | None:
    configured = os.environ.get("STARLANE_STATA_BIN")
    candidates = [
        configured,
        "/Applications/Stata/StataMP.app/Contents/MacOS/stata-mp",
        "/Applications/Stata/StataMP.app/Contents/MacOS/StataMP",
        "stata-mp",
        "stata-se",
        "stata",
    ]
    for candidate in candidates:
        if not candidate:
            continue
        if Path(candidate).exists() or shutil.which(candidate):
            return candidate
    return None


def require_columns(path: Path, columns: list[str]) -> pd.DataFrame:
    df = pd.read_csv(path)
    missing = [c for c in columns if c not in df.columns]
    if missing:
        raise AssertionError(f"Missing columns in {path}: {missing}")
    return df


def compare_summary_tables(stata_path: Path, python_path: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    stata = pd.read_csv(stata_path)
    python = pd.read_csv(python_path)
    common_keys = ["selection_id", "cv_idx", "vce_idx", "vce_suffix", "cv_selected"]
    for key in common_keys:
        assert key in stata.columns, f"Stata summary missing {key}"
        assert key in python.columns, f"Python summary missing {key}"
    assert set(stata["selection_id"]) == set(python["selection_id"])
    merged = stata[common_keys].merge(python[common_keys], on=common_keys, how="outer", indicator=True)
    assert (merged["_merge"] == "both").all(), merged[merged["_merge"] != "both"].to_string()
    return stata, python


def test_demo_enumeration_cardinality() -> None:
    args = RegressionArgs.from_list(SUMMARY_ARGS)
    cv_subsets = compute_cv_subsets(
        split_words(args.cv),
        split_words(args.cv_fixed),
        int(args.cv_min_count),
    )
    assert len(cv_subsets) == 31
    assert len(cv_subsets) * 4 == 124
    assert summary_chunk_path(PY_OUT).name == "combination_summary_part_0_1.csv"


def run_python_workflow(out_dir: Path) -> tuple[pd.DataFrame, Path]:
    env = {**os.environ, "STARLANE_EXPORT": str(out_dir), "STARLANE_TMP": str(out_dir / "tmp")}
    run([*UV_PYTHON, str(SCRIPTS / "regression_summary.py"), json.dumps(SUMMARY_ARGS, ensure_ascii=False)], env=env)
    summary_path = out_dir / "combination_summary.csv"
    summary = pd.read_csv(summary_path)
    first = summary.iloc[0]
    final_dir = out_dir / "final"
    final_args = [*SUMMARY_ARGS[:18], str(int(first["cv_idx"])), str(int(first["vce_idx"])), str(final_dir)]
    generated_source = final_dir / "regression_generated.py"
    run(
        [
            *UV_PYTHON,
            str(SCRIPTS / "generate_regression_py.py"),
            json.dumps(final_args, ensure_ascii=False),
            str(generated_source),
        ],
        env=env,
    )
    run([*UV_PYTHON, str(generated_source)], env=env)
    return summary, final_dir / "final_result.docx"


def write_stata_summary_runner(path: Path, export_dir: Path, tmp_dir: Path) -> None:
    quoted = " ".join(f'"{arg}"' for arg in SUMMARY_ARGS)
    path.write_text(
        "\n".join(
            [
                f'global STARLANE_EXPORT "{export_dir.as_posix()}"',
                f'global STARLANE_TMP "{tmp_dir.as_posix()}"',
                f'do "{(SCRIPTS / "regression_summary.do").as_posix()}" {quoted}',
                "exit",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def run_stata_batch(stata_bin: str, do_file: Path, cwd: Path) -> None:
    result = subprocess.run([stata_bin, "-b", "do", str(do_file)], cwd=cwd, text=True, capture_output=True)
    log_path = cwd / f"{do_file.stem}.log"
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr, file=sys.stderr)
        if log_path.exists():
            print(log_path.read_text(errors="replace")[-8000:], file=sys.stderr)
        raise SystemExit(result.returncode)
    if log_path.exists():
        log = log_path.read_text(errors="replace")
        if "r(" in log and "end of do-file" not in log:
            print(log[-8000:], file=sys.stderr)
            raise SystemExit("Stata log contains error")


def run_stata_workflow(stata_bin: str, out_dir: Path) -> tuple[pd.DataFrame, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    tmp_dir = out_dir / "tmp"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    runner = out_dir / "run_stata_summary.do"
    write_stata_summary_runner(runner, out_dir, tmp_dir)
    run_stata_batch(stata_bin, runner, out_dir)
    summary_path = summary_chunk_path(out_dir)
    if not summary_path.exists():
        raise AssertionError(f"Stata summary output not found: {summary_path}")
    canonical_summary_path = out_dir / "combination_summary.csv"
    if summary_path != canonical_summary_path:
        shutil.copyfile(summary_path, canonical_summary_path)
    summary = pd.read_csv(summary_path)
    first = summary.iloc[0]
    final_dir = out_dir / "final"
    final_dir.mkdir(parents=True, exist_ok=True)
    final_args = [*SUMMARY_ARGS[:18], str(int(first["cv_idx"])), str(int(first["vce_idx"])), str(final_dir)]
    generated_do = final_dir / "regression_generated.do"
    run(
        [
            *UV_PYTHON,
            str(SCRIPTS / "generate_regression_do.py"),
            json.dumps(final_args, ensure_ascii=False),
            str(generated_do),
        ]
    )
    run_stata_batch(stata_bin, generated_do, final_dir)
    docx = final_dir / "starlane-regression-results.docx"
    if not docx.exists():
        raise AssertionError(f"Stata final Word output not found: {docx}")
    return summary, docx


def test_python_backend() -> None:
    if not DEMO_DTA.exists():
        raise SystemExit(f"Missing demo dataset: {DEMO_DTA}")
    shutil.rmtree(OUT, ignore_errors=True)
    summary, python_docx = run_python_workflow(PY_OUT)

    summary_path = PY_OUT / "combination_summary.csv"
    expected_columns = [
        "selection_id",
        "cv_idx",
        "vce_idx",
        "vce_suffix",
        "cv_selected",
        "score",
        "baseline__lnApplyG__Attention__nocv",
        "baseline__lnGrantG__Attention__nocv",
        "baseline__lnApplyG__Attention__cv",
        "baseline__lnGrantG__Attention__cv",
        "robustness_alty__lnAGreenInv__Attention",
        "robustness_alty__lnGGreenInv__Attention",
        "robustness_lag__lnApplyG__Attention__l1",
        "robustness_lag__lnGrantG__Attention__l1",
        "iv__lnApplyG__Attention__Thermalinv__stage1",
        "iv__lnApplyG__Attention__Thermalinv__stage2",
        "iv__lnGrantG__Attention__Thermalinv__stage1",
        "iv__lnGrantG__Attention__Thermalinv__stage2",
        "mediation__Charge__lnApplyG__Attention",
        "mediation__Charge__M__Attention",
        "moderation__OverSea__lnApplyG__Attention",
    ]
    summary = require_columns(summary_path, expected_columns)
    assert summary.shape[0] == 8, summary.shape
    assert summary["selection_id"].is_unique

    final_dir = PY_OUT / "final"
    generated_source = final_dir / "regression_generated.py"

    final_path = final_dir / "final_result.csv"
    final = require_columns(final_path, ["column", "section", "depvar", "target", "coef", "se", "p_value", "nobs", "r2"])
    for column in expected_columns[6:]:
        assert column in set(final["column"]), column
    assert (final_dir / "final_result.md").exists()
    assert python_docx.exists()
    assert (final_dir / "python_backend_run_note.md").exists()
    assert (final_dir / "regression_generated.py").exists()
    doc = Document(final_dir / "final_result.docx")
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Starlane Python Regression Results" in text


def test_stata_python_comparison() -> None:
    stata_bin = find_stata()
    if not stata_bin:
        print("Stata binary not found; skipping Stata comparison.")
        return
    stata_summary, stata_docx = run_stata_workflow(stata_bin, STATA_OUT)
    python_summary = pd.read_csv(PY_OUT / "combination_summary.csv")
    compare_summary_tables(STATA_OUT / "combination_summary.csv", PY_OUT / "combination_summary.csv")
    assert stata_docx.exists()
    assert (PY_OUT / "final" / "final_result.docx").exists()
    stata_doc = Document(stata_docx)
    python_doc = Document(PY_OUT / "final" / "final_result.docx")
    assert len(stata_doc.tables) >= 1
    assert len(python_doc.tables) >= 1
    assert set(stata_summary["selection_id"]) == set(python_summary["selection_id"])


if __name__ == "__main__":
    test_demo_enumeration_cardinality()
    test_python_backend()
    test_stata_python_comparison()
    print("OK")
